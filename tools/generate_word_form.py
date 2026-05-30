#!/usr/bin/env python3
"""
生成启智归塾2026夏令营可交互Word报名表
支持内容控件（content controls），用户可在Word中直接填写
"""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '启智归塾2026夏令营报名表.docx')

def add_cc_text(doc, label, tag, placeholder="", required=False, hint=""):
    """添加纯文本内容控件"""
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    if required:
        r2 = p.add_run(' *')
        r2.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
        r2.font.size = Pt(11)
    if hint:
        p2 = doc.add_paragraph()
        r3 = p2.add_run(hint)
        r3.font.size = Pt(9)
        r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        r3.italic = True
    # Add content control
    sdt = OxmlElement('w:sdt')
    sdtPr = OxmlElement('w:sdtPr')
    alias = OxmlElement('w:alias')
    alias.set(qn('w:val'), tag)
    sdtPr.append(alias)
    tag_el = OxmlElement('w:tag')
    tag_el.set(qn('w:val'), tag)
    sdtPr.append(tag_el)
    placeholder_el = OxmlElement('w:placeholder')
    docPart = OxmlElement('w:docPart')
    docPart.set(qn('w:val'), placeholder)
    placeholder_el.append(docPart)
    sdtPr.append(placeholder_el)
    sdt.append(sdtPr)
    sdtContent = OxmlElement('w:sdtContent')
    p_cc = OxmlElement('w:p')
    r_cc = OxmlElement('w:r')
    t_cc = OxmlElement('w:t')
    t_cc.text = placeholder
    r_cc.append(t_cc)
    p_cc.append(r_cc)
    sdtContent.append(p_cc)
    sdt.append(sdtContent)
    doc.element.body.append(sdt)
    return sdt

def add_cc_dropdown(doc, label, tag, options, required=False):
    """添加下拉列表内容控件"""
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    if required:
        r2 = p.add_run(' *')
        r2.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
        r2.font.size = Pt(11)
    sdt = OxmlElement('w:sdt')
    sdtPr = OxmlElement('w:sdtPr')
    tag_el = OxmlElement('w:tag')
    tag_el.set(qn('w:val'), tag)
    sdtPr.append(tag_el)
    dd = OxmlElement('w:dropDownList')
    for i, opt in enumerate(options):
        li = OxmlElement('w:listItem')
        li.set(qn('w:displayText'), opt)
        li.set(qn('w:value'), opt)
        dd.append(li)
    sdtPr.append(dd)
    sdt.append(sdtPr)
    sdtContent = OxmlElement('w:sdtContent')
    p_cc = OxmlElement('w:p')
    r_cc = OxmlElement('w:r')
    t_cc = OxmlElement('w:t')
    t_cc.text = options[0] if options else ''
    r_cc.append(t_cc)
    p_cc.append(r_cc)
    sdtContent.append(p_cc)
    sdt.append(sdtContent)
    doc.element.body.append(sdt)
    return sdt

def add_cc_checkbox(doc, label, tag):
    """添加复选框"""
    sdt = OxmlElement('w:sdt')
    sdtPr = OxmlElement('w:sdtPr')
    tag_el = OxmlElement('w:tag')
    tag_el.set(qn('w:val'), tag)
    sdtPr.append(tag_el)
    cb = OxmlElement('w:checkbox')
    checked = OxmlElement('w:checked')
    checked.set(qn('w:val'), '0')
    cb.append(checked)
    sdtPr.append(cb)
    sdt.append(sdtPr)
    sdtContent = OxmlElement('w:sdtContent')
    p_cc = OxmlElement('w:p')
    r_cc = OxmlElement('w:r')
    t_cc = OxmlElement('w:t')
    t_cc.text = ' ' + label
    r_cc.append(t_cc)
    p_cc.append(r_cc)
    sdtContent.append(p_cc)
    sdt.append(sdtContent)
    doc.element.body.append(sdt)
    return sdt

def add_heading(doc, text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.name = '微软雅黑'
    return h

def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '5B7B5E')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_form():
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Title
    title = doc.add_heading('启智归塾 · 2026夏令营报名表', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('启智书院 × 日月洲度假村 × 江淮书院  联合主办')
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x5B, 0x7B, 0x5E)
    
    doc.add_paragraph()  # spacer
    
    # ── Section 1: 联系人信息 ──
    add_heading(doc, '一、联系人信息')
    add_cc_text(doc, '联系人姓名', 'parent_name', '请输入您的姓名', required=True)
    add_cc_text(doc, '手机号', 'phone', '请输入11位手机号', required=True, hint='（必填，用于后续联系确认）')
    add_cc_text(doc, '微信号', 'wechat', '请输入微信号（选填）', hint='（选填，便于微信联系）')
    doc.add_paragraph()
    
    # ── Section 2: 孩子信息 ──
    add_heading(doc, '二、孩子信息')
    for i in range(1, 4):
        p = doc.add_paragraph()
        r = p.add_run(f'孩子{i}')
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = '微软雅黑'
        if i == 1:
            r2 = p.add_run('（至少填一个，可填2-3个孩子）')
            r2.font.size = Pt(9)
            r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        
        add_cc_text(doc, f'  姓名', f'child{i}_name', f'孩子{i}姓名', required=(i==1))
        add_cc_dropdown(doc, f'  性别', f'child{i}_gender', ['', '男', '女'], required=(i==1))
        add_cc_text(doc, f'  年龄(5-18)', f'child{i}_age', '如：10', required=(i==1))
        add_cc_dropdown(doc, f'  年级', f'child{i}_grade', ['', '学前', '小1-3', '小4-6', '初中', '高中'])
        add_cc_dropdown(doc, f'  特殊需求', f'child{i}_special', ['无', '有'])
        add_cc_text(doc, f'  特殊需求说明', f'child{i}_special_detail', '如：食物过敏、药物等')
        doc.add_paragraph()
    
    # ── Section 3: 报名选型 ──
    add_heading(doc, '三、报名选型')
    add_cc_dropdown(doc, '选择产品', 'product', ['', '体验版(7天)2980元', '进阶版(14天)4980元', '完整版(21天)6980元'], required=True)
    doc.add_paragraph()
    
    # ── Section 4: 家长陪同 ──
    add_heading(doc, '四、家长陪同（选填）')
    add_cc_dropdown(doc, '母亲', 'mother_accompany', ['不参加', '全程 ¥3,580', '按周 ¥980/7天'])
    p_m = doc.add_paragraph()
    r_m = p_m.add_run('母亲选择周次（按周时勾选）：')
    r_m.font.size = Pt(10)
    r_m.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    add_cc_checkbox(doc, '第一周（8月1日-7日）', 'mother_week1')
    add_cc_checkbox(doc, '第二周（8月8日-14日）', 'mother_week2')
    add_cc_checkbox(doc, '第三周（8月15日-21日）', 'mother_week3')
    
    doc.add_paragraph()
    add_cc_dropdown(doc, '父亲', 'father_accompany', ['不参加', '全程 ¥3,580', '按周 ¥980/7天'])
    p_f = doc.add_paragraph()
    r_f = p_f.add_run('父亲选择周次（按周时勾选）：')
    r_f.font.size = Pt(10)
    r_f.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    add_cc_checkbox(doc, '第一周（8月1日-7日）', 'father_week1')
    add_cc_checkbox(doc, '第二周（8月8日-14日）', 'father_week2')
    add_cc_checkbox(doc, '第三周（8月15日-21日）', 'father_week3')
    doc.add_paragraph()
    
    # ── Section 5: 开放性问题 ──
    add_heading(doc, '五、开放性问题')
    add_cc_text(doc, '1. 孩子最近一次让你意外或触动的事是什么？', 'qa1', '请在这里写下您的回答', required=True)
    add_cc_text(doc, '2. 您对这次夏令营的期待是什么？', 'qa2', '请在这里写下您的回答', required=True)
    doc.add_paragraph()
    
    # ── Section 6: 其他 ──
    add_heading(doc, '六、其他信息（选填）')
    add_cc_text(doc, '推荐人', 'referrer', '如有人推荐请填写')
    add_cc_dropdown(doc, '获知渠道', 'source', ['', '微信朋友圈', '公众号', '朋友推荐', '抖音/视频号', '其他'])
    add_cc_text(doc, '备注', 'notes', '如有特殊要求请在此说明')
    doc.add_paragraph()
    
    # Footer note
    p_note = doc.add_paragraph()
    r_note = p_note.add_run('填写完毕后请保存，通过微信发回给启智书院工作人员。')
    r_note.font.size = Pt(9)
    r_note.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    r_note.italic = True
    
    doc.save(OUTPUT)
    print(f'✅ Word报名表已生成: {OUTPUT}')

if __name__ == '__main__':
    create_form()
