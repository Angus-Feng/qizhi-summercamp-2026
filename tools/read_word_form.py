#!/usr/bin/env python3
"""
读取启智归塾2026夏令营报名表Word文件，提取数据并提交到报名数据库

用法:
  python3 tools/read_word_form.py 张三填好的报名表.docx
  python3 tools/read_word_form.py 张三填好的报名表.docx --dry-run   # 仅预览，不提交
  python3 tools/read_word_form.py 张三填好的报名表.docx --local      # 提交到本地服务器

文件路径建议放到 summercamp/word_submissions/ 目录下
"""
import sys, os, re, json, argparse
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..'))

from docx import Document

# ── API 地址 ────────────────────────────────
API_URL = 'https://summercamp.qizhishuyuan.cn/api/register'
LOCAL_URL = 'http://localhost:3000/api/register'

# ── 价格常量 ────────────────────────────────
PRODUCT_PRICES = { '7': 2980, '14': 4980, '21': 6980 }
ACCOMPANY_RATE_FULL = 3580
ACCOMPANY_RATE_7DAY = 980

def extract_table_data(doc):
    """从Word表格中提取所有单元格文本，按表分组"""
    result = []
    for table in doc.tables:
        rows_data = []
        for row in table.rows:
            cells_text = []
            for cell in row.cells:
                text = cell.text.strip()
                cells_text.append(text)
            rows_data.append(cells_text)
        result.append(rows_data)
    return result

def extract_form_data(doc):
    """根据表单结构提取字段（勾选版）"""
    data = {}
    tables = extract_table_data(doc)

    def cell_vals(tidx):
        rows = tables[tidx] if tidx < len(tables) else []
        vals = []
        for row in rows:
            if len(row) > 1:
                v = row[1].strip()
                if v and v != 'None' and '请选择' not in v and '请输入' not in v:
                    vals.append(v)
                else:
                    vals.append('')
        return vals

    def is_checked(para_text, opt):
        """检查段落文本中某选项是否被勾选"""
        idx = para_text.find(opt)
        if idx < 0: return False
        pre = para_text[max(0,idx-3):idx]
        if any(c in pre for c in ['☑','✓','✔','●','■']): return True
        if '□'+opt in para_text or '□ '+opt in para_text: return False
        return True  # □被删 = 视为已选

    def find_para_after(keyword, start=0):
        """在段落中查找包含某关键词的第一个段落索引"""
        for i, p in enumerate(doc.paragraphs):
            if i < start: continue
            if keyword in p.text:
                return i, p.text
        return -1, ''

    def parse_accompany(text):
        if is_checked(text, '按周'): return 'weekly'
        if is_checked(text, '全程'): return 'full'
        return 'no'

    # ── 表0: 联系人 ──
    if len(tables) >= 1:
        v = cell_vals(0)
        if len(v)>=1: data['parent_name'] = v[0]
        if len(v)>=2: data['relation'] = v[1]
        if len(v)>=3: data['phone'] = v[2]
        if len(v)>=4: data['wechat'] = v[3]

    # ── 孩子(表1/3/5=基本信息, 表2/4/6=特殊说明) + 段落勾选 ──
    children = []
    for ci in range(3):
        base_tbl = 1 + ci*2  # 1, 3, 5
        detail_tbl = 2 + ci*2  # 2, 4, 6
        if base_tbl >= len(tables): break
        v = cell_vals(base_tbl)
        if len(v)<1 or not v[0] or '孩子' in v[0]: continue
        child = {}
        if len(v)>=1: child['name'] = v[0]
        if len(v)>=2: child['gender'] = v[1] if v[1] in ('男','女') else ''
        if len(v)>=3:
            try: child['age'] = int(re.sub(r'[^\d]','',v[2]))
            except: child['age'] = 0
        if len(v)>=4: child['id_number'] = v[3]
        child['has_special_needs'] = 'no'
        child['grade'] = ''
        if detail_tbl < len(tables):
            dv = cell_vals(detail_tbl)
            if len(dv)>=1: child['special_needs_detail'] = dv[0]
        children.append(child)

    # ── 段落勾选：年级 ──
    GRADE_OPTS = ['学前','小1-3','小4-6','初中','高中']
    ci = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if '年级' in t and '勾选' in t:
            if ci < len(children):
                for g in GRADE_OPTS:
                    if is_checked(t, g):
                        children[ci]['grade'] = g
                        break
                ci += 1

    # ── 段落勾选：特殊需求 ──
    ci = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if '特殊需求' in t and '勾选' in t:
            if ci < len(children):
                if is_checked(t, '有') and not is_checked(t, '如有'):
                    children[ci]['has_special_needs'] = 'yes'
                ci += 1

    data['children'] = children

    # ── 产品勾选 ──
    data['product'] = ''
    for p in doc.paragraphs:
        t = p.text.strip()
        if is_checked(t, '体验版'): data['product'] = '7'; break
        if is_checked(t, '进阶版'): data['product'] = '14'; break
        if is_checked(t, '完整版'): data['product'] = '21'; break

    # ── 家长陪同 + 周次（段落勾选） ──
    accomp_state = {'mother':'no','father':'no','other':'no'}
    weeks_state = {'mother':[],'father':[],'other':[]}
    current = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if '母亲陪同' in t: current = 'mother'
        elif '父亲陪同' in t: current = 'father'
        elif '其它亲属' in t: current = 'other'
        elif current and '参与方式' in t:
            accomp_state[current] = parse_accompany(t)
        elif current and '第一周' in t:
            if is_checked(t, '第一周'): weeks_state[current].append(1)
        elif current and '第二周' in t:
            if is_checked(t, '第二周'): weeks_state[current].append(2)
        elif current and '第三周' in t:
            if is_checked(t, '第三周'): weeks_state[current].append(3)
        elif current and ('母亲' in t or '父亲' in t or '其它亲属' in t) and current not in t:
            current = None

    data['mother_accompany'] = accomp_state['mother']
    data['father_accompany'] = accomp_state['father']
    data['mother_weeks'] = weeks_state['mother'] if accomp_state['mother']=='weekly' else []
    data['father_weeks'] = weeks_state['father'] if accomp_state['father']=='weekly' else []

    # ── 其它亲属 表7: 关系 ──
    data['other_accompany'] = accomp_state['other']
    data['other_relation'] = ''
    data['other_weeks'] = weeks_state['other'] if accomp_state['other']=='weekly' else []
    if len(tables) >= 8:
        v = cell_vals(7)
        if len(v)>=1: data['other_relation'] = v[0]

    # ── QAs 表8 ──
    if len(tables) >= 9:
        v = cell_vals(8)
        if len(v)>=1: data['qa1'] = v[0]
        if len(v)>=2: data['qa2'] = v[1]

    # ── 其他 表9: referrer, notes ──
    if len(tables) >= 10:
        v = cell_vals(9)
        if len(v)>=1: data['referrer'] = v[0] if '推荐' not in v[0] else ''
        if len(v)>=2: data['notes'] = v[1] if '特殊' not in v[1] else ''

    # ── 获知渠道 段落勾选 ──
    data['source'] = ''
    for p in doc.paragraphs:
        t = p.text.strip()
        if '获知渠道' in t:
            if is_checked(t, '公众号'): data['source'] = '公众号'; break
            if is_checked(t, '朋友推荐'): data['source'] = '朋友推荐'; break
            if is_checked(t, '抖音'): data['source'] = '抖音/视频号'; break
            if is_checked(t, '微信朋友圈'): data['source'] = '微信朋友圈'; break
            if is_checked(t, '其他'): data['source'] = '其他'; break

    # 默认值
    for k in ['parent_name','phone','wechat','product','qa1','qa2','referrer','source','notes']:
        data.setdefault(k,'')
    for k in ['father_accompany','mother_accompany','other_accompany']:
        data.setdefault(k,'no')

    # 价格计算
    product = data['product']
    base_price = PRODUCT_PRICES.get(product, 0)
    data['child_count'] = len(data['children'])
    data['base_price'] = base_price
    data['children_total'] = data['child_count'] * base_price
    accompany_fee = 0
    for pfx in ['father','mother','other']:
        if data.get(f'{pfx}_accompany') == 'full': accompany_fee += ACCOMPANY_RATE_FULL
        elif data.get(f'{pfx}_accompany') == 'weekly': accompany_fee += len(data.get(f'{pfx}_weeks',[])) * ACCOMPANY_RATE_7DAY
    data['accompany_fee'] = accompany_fee
    data['total_price'] = data['children_total'] + accompany_fee
    return data

def validate_data(data):
    """校验数据"""
    errors = []
    if not data.get('parent_name'): errors.append('联系人姓名不能为空')
    phone = data.get('phone', '')
    if not re.match(r'^1[3-9]\d{9}$', phone): errors.append('请输入正确的11位手机号')
    if not data.get('children'):
        errors.append('请至少填写一个孩子')
    else:
        for i, ch in enumerate(data['children']):
            if not ch.get('name'): errors.append(f'孩子{i+1}姓名不能为空')
            if not ch.get('gender') or ch['gender'] not in ('男', '女'): errors.append(f'孩子{i+1}性别无效')
            if not ch.get('age') or ch['age'] < 5 or ch['age'] > 18: errors.append(f'孩子{i+1}年龄需在5-18岁')
            if not ch.get('id_number') or not re.match(r'^\d{17}[\dXx]$', ch.get('id_number', '')): errors.append(f'孩子{i+1}身份证号格式错误')
    if not data.get('product') or data['product'] not in ('7', '14', '21'): errors.append('请选择报名产品')
    if not data.get('qa1'): errors.append('请填写开放性问题1')
    if not data.get('qa2'): errors.append('请填写开放性问题2')
    return errors

def print_summary(data):
    """打印数据摘要"""
    print('\n📋 报名数据预览：')
    print(f'  联系人：{data["parent_name"]} | {data["phone"]}')
    if data.get('wechat'): print(f'  微信：{data["wechat"]}')
    print(f'  孩子数：{data["child_count"]}')
    for i, ch in enumerate(data['children']):
        print(f'    孩子{i+1}：{ch.get("name","?")} | {ch.get("gender","?")} | {ch.get("age","?")}岁 | {ch.get("grade","")}')
    p_names = {'7':'体验版(7天)','14':'进阶版(14天)','21':'完整版(21天)'}
    print(f'  产品：{p_names.get(data["product"],data["product"])}')
    print(f'  母亲：{data["mother_accompany"]} 周次={data.get("mother_weeks",[])}')
    print(f'  父亲：{data["father_accompany"]} 周次={data.get("father_weeks",[])}')
    print(f'  Q1：{data.get("qa1","")[:50]}...')
    print(f'  Q2：{data.get("qa2","")[:50]}...')
    print(f'  合计：¥{data["total_price"]:,}')

def submit_data(data, api_url):
    """提交到API"""
    import urllib.request
    try:
        body = json.dumps(data, ensure_ascii=False).encode('utf-8')
        req = urllib.request.Request(api_url, data=body, headers={'Content-Type': 'application/json'}, method='POST')
        with urllib.request.urlopen(req, timeout=30) as resp:
            result = json.loads(resp.read().decode('utf-8'))
            return result
    except urllib.error.HTTPError as e:
        err_body = e.read().decode('utf-8')
        try:
            return json.loads(err_body)
        except:
            return {'success': False, 'errors': [f'HTTP {e.code}: {err_body}']}
    except Exception as e:
        return {'success': False, 'errors': [str(e)]}

def main():
    parser = argparse.ArgumentParser(description='读取Word报名表并提交到报名数据库')
    parser.add_argument('file', help='Word报名表文件路径(.docx)')
    parser.add_argument('--dry-run', action='store_true', help='仅预览不提交')
    parser.add_argument('--local', action='store_true', help='提交到本地服务器')
    parser.add_argument('--url', help='自定义API地址')
    args = parser.parse_args()
    
    if not os.path.exists(args.file):
        print(f'❌ 文件不存在: {args.file}')
        sys.exit(1)
    
    print(f'📖 读取文件: {args.file}')
    doc = Document(args.file)
    data = extract_form_data(doc)
    
    print_summary(data)
    
    errors = validate_data(data)
    if errors:
        print(f'\n❌ 数据校验失败：')
        for e in errors:
            print(f'  - {e}')
        if not args.dry_run:
            print('\n请修正Word文件后重试，或用 --dry-run 先预览')
            sys.exit(1)
        else:
            print('\n（dry-run模式，不提交）')
            return
    
    if args.dry_run:
        print('\n✅ 数据校验通过（dry-run模式，未提交）')
        return
    
    api_url = args.url or (LOCAL_URL if args.local else API_URL)
    print(f'\n📤 提交到: {api_url}')
    result = submit_data(data, api_url)
    
    if result.get('success'):
        print(f'✅ 提交成功！报名ID: {result.get("data",{}).get("id","?")}')
    else:
        print(f'❌ 提交失败：')
        for e in result.get('errors', ['未知错误']):
            print(f'  - {e}')

if __name__ == '__main__':
    main()
